#!/usr/bin/env python3
"""
ISARVA HRMS vs Karnataka HRMS 2.0 — Gap Analysis Document Generator

Generates:
  - HRMS_GAP_ANALYSIS_KARNATAKA_2.0.md
  - HRMS_GAP_ANALYSIS_KARNATAKA_2.0.pdf
  - HRMS_GAP_ANALYSIS_KARNATAKA_2.0.docx

Usage:
  python3 generate_gap_analysis.py
"""

from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUTPUT_DIR = Path(__file__).resolve().parent
TODAY = date.today().strftime("%d %B %Y")

META = {
    "title": "HRMS Gap Analysis: ISARVA vs Karnataka HRMS 2.0",
    "subtitle": "Payroll & Attendance Modules — Deep Comparison",
    "prepared_for": "ISARVA HRMS (Payroll & Attendance)",
    "benchmark": "Karnataka State Government HRMS 2.0 (hrms.karnataka.gov.in / HRMS2 ESS)",
    "version": "1.0",
    "date": TODAY,
}

STATUS = {
    "full": ("Available", "FULL", colors.HexColor("#1B7F3A")),
    "partial": ("Partially Available", "PARTIAL", colors.HexColor("#B8860B")),
    "missing": ("Not Available", "GAP", colors.HexColor("#C0392B")),
    "na": ("Not Applicable", "N/A", colors.HexColor("#5D6D7E")),
}

# ---------------------------------------------------------------------------
# Gap analysis data — each module contains feature rows
# status: full | partial | missing | na
# ---------------------------------------------------------------------------

MODULES: list[dict] = [
    {
        "id": "1",
        "title": "Executive Summary",
        "intro": (
            "Karnataka HRMS 2.0 is a state-government platform built for 6+ lakh government "
            "employees, aided institutions, and university staff. It is designed around government "
            "pay rules (pay scales, GPF, NPS, DA), DDO-led salary processing, and deep integration "
            "with 45+ departments.\n\n"
            "ISARVA HRMS (Payroll + Attendance) is a private-sector, multi-tenant SaaS platform "
            "focused on CTC-based salary, EPF/ESI compliance, monthly payroll wizard, and "
            "attendance-to-payroll integration for companies.\n\n"
            "This is not a like-for-like comparison of two identical products. Karnataka HRMS 2.0 "
            "serves government HR lifecycle; ISARVA serves corporate payroll operations. Many "
            "government-only features are intentionally absent. Where both overlap (payroll, leave, "
            "attendance, ESS), this document highlights what ISARVA already does well and what "
            "would need to be built to match Karnataka HRMS 2.0 depth."
        ),
        "features": [
            {
                "name": "Overall platform maturity",
                "karnataka": "20+ years evolution (HRMS 1.0 since 2005); HRMS 2.0 with Kellton; 200+ APIs; 6 lakh users",
                "isarva": "Modern Laravel 12 multi-app architecture; multi-tenant; Payroll + Attendance + Hub SSO",
                "status": "partial",
                "gap": "ISARVA has modern tech stack but far fewer government integrations and smaller proven scale at state level.",
                "recommendation": "Position ISARVA for private sector; do not try to replicate full govt HRMS without dedicated govt product line.",
            },
            {
                "name": "Target user base",
                "karnataka": "Government, aided schools, boards, universities — KGID-based identity",
                "isarva": "Private companies, SMEs, multi-location corporates — company code + email login",
                "status": "na",
                "gap": "Different markets by design.",
                "recommendation": "No action unless entering government sector.",
            },
        ],
    },
    {
        "id": "2",
        "title": "Authentication, Access & Security",
        "intro": (
            "Karnataka HRMS uses KGID (employee insurance ID) as the primary key, with separate "
            "login portals for Government, Aided, and Boards/Universities. The new ESS portal "
            "(hrmsess.karnataka.gov.in) uses OTP on registered mobile. DDOs (Drawing and "
            "Disbursing Officers) register employees — no self-registration."
        ),
        "features": [
            {
                "name": "Employee login (KGID / Employee Code)",
                "karnataka": "KGID number + password on main portal; KGID + OTP on ESS mobile/web",
                "isarva": "Company code + email/employee ID + password via Workspace Hub; JWT SSO to Payroll/Attendance",
                "status": "partial",
                "gap": "No KGID concept; no OTP-only ESS login; corporate SSO model instead.",
                "recommendation": "Add optional OTP login for employee self-service if customers request it.",
            },
            {
                "name": "Three-tier login categories (Govt / Aided / Boards)",
                "karnataka": "Separate login tabs and workflows per employee category",
                "isarva": "Single tenant type per company; no aided-institution or university category split",
                "status": "missing",
                "gap": "No equivalent employee-category portals.",
                "recommendation": "Only needed for government deployments.",
            },
            {
                "name": "DDO (salary disbursing officer) role",
                "karnataka": "21,000+ DDOs process monthly salary accounts for their establishment",
                "isarva": "Payroll admin / HR admin with role-based permissions; no DDO establishment hierarchy",
                "status": "partial",
                "gap": "No DDO code, establishment section, or govt salary bill workflow.",
                "recommendation": "Map 'Payroll Admin' as functional equivalent for private sector.",
            },
            {
                "name": "Self-registration",
                "karnataka": "Not allowed — DDO registers employee and sends initial password to mobile",
                "isarva": "Admin creates users; optional Google OAuth; employee cannot self-register",
                "status": "full",
                "gap": "Aligned — admin-controlled registration.",
                "recommendation": "None.",
            },
            {
                "name": "Password management",
                "karnataka": "Change password with DDO code + role selection; ESS uses OTP",
                "isarva": "Forgot password, profile password change, cross-app password sync with Attendance",
                "status": "full",
                "gap": "Corporate password flows are in place.",
                "recommendation": "None.",
            },
            {
                "name": "Aadhaar e-KYC integration",
                "karnataka": "Kartavya-KAAMS attendance app integrates Aadhaar e-KYC for identity verification",
                "isarva": "Aadhaar number stored as employee personal field only — no e-KYC verification",
                "status": "missing",
                "gap": "No Aadhaar-based identity verification at login or attendance.",
                "recommendation": "Integrate Aadhaar e-KYC API if attendance fraud prevention is a priority.",
            },
            {
                "name": "Multi-tenant / company isolation",
                "karnataka": "Department-level segregation within single govt system",
                "isarva": "Full multi-tenant: central registry + per-tenant database shards",
                "status": "full",
                "gap": "ISARVA exceeds Karnataka model for SaaS multi-company hosting.",
                "recommendation": "Strength to highlight for private sector clients.",
            },
        ],
    },
    {
        "id": "3",
        "title": "Employee Master Data & Service Book",
        "intro": (
            "Karnataka HRMS maintains an Electronic Service Book (e-SR) — a complete career record "
            "including every posting, promotion, transfer, training, and suspension. This is the "
            "heart of government HR. ISARVA maintains employee master data suitable for payroll "
            "but not a lifelong government service register."
        ),
        "features": [
            {
                "name": "Electronic Service Book (e-SR)",
                "karnataka": "Full career history: postings, promotions, transfers, training, suspensions — viewable by employee via ESS",
                "isarva": "Employee profile with personal, bank, documents, salary history; increment records; exit details",
                "status": "missing",
                "gap": "No service book with posting history, cadre details, or govt order linkage.",
                "recommendation": "Build 'Employee Timeline' module if long-tenure career tracking is needed.",
            },
            {
                "name": "KGID / Government ID linkage",
                "karnataka": "KGID links employee to insurance, ESS login, and service records",
                "isarva": "Employee ID, UAN (EPF), PAN, Aadhaar fields on employee profile",
                "status": "partial",
                "gap": "Private-sector IDs present; no government KGID or PRAN lifecycle management.",
                "recommendation": "Add PRAN/NPS fields if serving post-2004 govt-style pension clients.",
            },
            {
                "name": "Pay Scale & Basic Pay (Government)",
                "karnataka": "Pay scale, basic pay, DA%, HRA%, grade pay — revised per pay commission",
                "isarva": "CTC-based monthly/annual salary with configurable earning components (Basic, HRA, etc.)",
                "status": "partial",
                "gap": "No pay commission matrix, pay level, or automatic DA revision per govt notification.",
                "recommendation": "For govt clients: build pay matrix module. For private: current CTC model is correct.",
            },
            {
                "name": "Cadre management",
                "karnataka": "Cadre, category, class of post — drives promotion and transfer eligibility",
                "isarva": "Department, designation, location, employee status — no cadre hierarchy",
                "status": "missing",
                "gap": "No cadre/category/class-of-post master or rules engine.",
                "recommendation": "Government-sector feature only.",
            },
            {
                "name": "Employee documents",
                "karnataka": "Transfer orders, posting orders accessible via ESS",
                "isarva": "Document types, file upload on employee profile, resignation letter PDF",
                "status": "partial",
                "gap": "Document storage exists but no govt order workflow or transfer order generation.",
                "recommendation": "Extend document module with order templates if needed.",
            },
            {
                "name": "Suspension management",
                "karnataka": "Suspension details with subsistence allowance calculation in salary",
                "isarva": "Employee status (active/left); held salary module; no subsistence allowance rules",
                "status": "missing",
                "gap": "No suspension workflow or subsistence pay calculation.",
                "recommendation": "Add only if serving government or PSU clients.",
            },
            {
                "name": "Profile update by employee",
                "karnataka": "Limited self-update via ESS; major changes via DDO",
                "isarva": "Profile view in Attendance; data change request (DPDP compliance) sent to Payroll admin",
                "status": "partial",
                "gap": "DPDP data change request exists — good for private sector; not full ESS profile edit.",
                "recommendation": "Expand self-service profile fields (contact, emergency, bank) with approval workflow.",
            },
        ],
    },
    {
        "id": "4",
        "title": "Recruitment, Transfer & Promotion",
        "intro": (
            "Karnataka HRMS 2.0 (Kellton scope) includes recruitment, cadre management, transfers, "
            "and promotion as first-class modules with API integration across 45+ departments. "
            "ISARVA handles increments and exit but not full recruitment or transfer lifecycle."
        ),
        "features": [
            {
                "name": "Recruitment module",
                "karnataka": "End-to-end recruitment integrated with HRMS 2.0 (Kellton project scope)",
                "isarva": "Not in Payroll or Attendance; no job posting, application, or selection workflow",
                "status": "missing",
                "gap": "No recruitment / ATS module.",
                "recommendation": "Separate product module or integrate with third-party ATS.",
            },
            {
                "name": "Transfer management",
                "karnataka": "Transfer orders, posting changes, inter-department transfers; viewable on ESS",
                "isarva": "Location and department can be edited on employee; no transfer order workflow",
                "status": "missing",
                "gap": "No transfer request, approval chain, relieving/joining, or order document.",
                "recommendation": "Build transfer workflow if enterprise clients need internal mobility tracking.",
            },
            {
                "name": "Promotion management",
                "karnataka": "Promotion history in service book; linked to pay scale revision",
                "isarva": "Increment & Promotion module (IncrementController) — salary revision with history and revert",
                "status": "partial",
                "gap": "Salary increment exists; no designation promotion workflow or pay scale upgrade.",
                "recommendation": "Link increment module to designation change and approval chain.",
            },
            {
                "name": "Annual Performance Appraisal (APR)",
                "karnataka": "Annual performance management module (Kellton HRMS 2.0 scope)",
                "isarva": "Not available in Payroll or Attendance",
                "status": "missing",
                "gap": "No performance review, rating, or APR-linked increment.",
                "recommendation": "Separate PMS module; optional link to increment.",
            },
        ],
    },
    {
        "id": "5",
        "title": "Payroll Processing & Salary Disbursement",
        "intro": (
            "Both systems process monthly salary, but the models differ fundamentally. Karnataka uses "
            "government pay components (Basic, DA, HRA, allowances, GPF deduction). ISARVA uses "
            "CTC-based components prorated by attendance days."
        ),
        "features": [
            {
                "name": "Monthly salary processing",
                "karnataka": "DDO submits salary bill monthly; processed centrally; 6 lakh+ accounts",
                "isarva": "5-step payroll wizard: select month → attendance → salary review → compare → finalize",
                "status": "full",
                "gap": "Core monthly payroll exists with lock/finalize.",
                "recommendation": "None for private sector.",
            },
            {
                "name": "Attendance-based proration",
                "karnataka": "Present days / LOP integrated with leave and attendance systems",
                "isarva": "worked_days / total_working_days factor applied to all earnings and deductions",
                "status": "full",
                "gap": "Strong integration — attendance locked in Attendance app, consumed by Payroll.",
                "recommendation": "Highlight as competitive strength.",
            },
            {
                "name": "Dearness Allowance (DA) auto-revision",
                "karnataka": "DA revised by government notification; applied across all employees automatically",
                "isarva": "DA can be a salary component but no govt notification-driven mass revision",
                "status": "partial",
                "gap": "No DA rate master linked to govt gazette notifications.",
                "recommendation": "Add bulk component revision tool for DA/rate changes.",
            },
            {
                "name": "Salary slip (payslip)",
                "karnataka": "PDF download via portal or ESS app; includes Basic, DA, HRA, GPF, GIS, PT, recoveries",
                "isarva": "PDF payslip generation, email, bulk send; employee views via Attendance portal API",
                "status": "full",
                "gap": "Payslip delivery is implemented.",
                "recommendation": "Add annual salary statement (see Tax module).",
            },
            {
                "name": "Annual salary statement",
                "karnataka": "Full-year summary for income tax filing — available on ESS",
                "isarva": "Monthly analytics and year comparison reports; no dedicated Form 16 / annual tax statement",
                "status": "partial",
                "gap": "No employee-facing annual tax summary document.",
                "recommendation": "Generate annual earnings/deductions PDF for IT filing.",
            },
            {
                "name": "Bank salary transfer file",
                "karnataka": "Integrated with treasury / government banking systems",
                "isarva": "Bank transfer Excel, CSV, ICICI XLSX formats with NEFT/RTGS/IMPS types",
                "status": "full",
                "gap": "Private-sector bank file export is strong.",
                "recommendation": "Add more bank format templates as needed.",
            },
            {
                "name": "Multi-location payroll",
                "karnataka": "Department and establishment level processing",
                "isarva": "Per-location or global payroll run with location_id on payout month status",
                "status": "full",
                "gap": "Multi-location payroll supported.",
                "recommendation": "None.",
            },
            {
                "name": "Held salary / early salary",
                "karnataka": "Recoveries and advance adjustments in salary bill",
                "isarva": "Held salary module, salary advances with auto-deduction at finalize, early salary flag",
                "status": "full",
                "gap": "Advance and hold mechanisms exist.",
                "recommendation": "None.",
            },
            {
                "name": "OT, incentives, holiday payout",
                "karnataka": "Allowances processed as part of govt salary components",
                "isarva": "Dedicated OT/Incentive module with separate finalize flags; holiday work payout",
                "status": "full",
                "gap": "OT and incentives well covered for private sector.",
                "recommendation": "None.",
            },
            {
                "name": "Full & Final Settlement (FFS) on exit",
                "karnataka": "Retirement benefits and final settlement via pension/GPF modules",
                "isarva": "Exit employee workflow (Pending → Approved → Completed); FFS calculation",
                "status": "partial",
                "gap": "Exit exists; no gratuity/pension settlement automation.",
                "recommendation": "Add gratuity calculation on exit for private sector compliance.",
            },
        ],
    },
    {
        "id": "6",
        "title": "Tax & Statutory Compliance — Deep Analysis",
        "intro": (
            "This is the most important area of difference. Karnataka HRMS 2.0 follows government "
            "statutory rules (GPF, NPS, GIS, KGID, DA-linked PT). ISARVA follows private-sector "
            "Indian payroll compliance (EPF, ESI, simplified PT, optional TDS slabs in settings).\n\n"
            "Plain-language summary: ISARVA is built for companies under the Companies Act / "
            "Factories Act with EPF and ESI. Karnataka HRMS is built for government employees "
            "under GPF/NPS rules. They share the word 'payroll' but use different tax and "
            "deduction engines."
        ),
        "features": [
            {
                "name": "EPF (Employees Provident Fund) — Private Sector",
                "karnataka": "Not applicable — government uses GPF instead of EPF",
                "isarva": "Full EPF engine: 12% employee share, optional 24% if employer share deducted from CTC, wage cap ₹15,000 or full wages, manual override, UAN on employee, ECR Excel export",
                "status": "full",
                "gap": "ISARVA is strong for private-sector EPF — Karnataka govt does not use this.",
                "recommendation": "Core strength for corporate clients.",
            },
            {
                "name": "GPF (General Provident Fund) — Government",
                "karnataka": "GPF account with balance, passbook, loan status, installment deduction in salary",
                "isarva": "Not available — no GPF account, passbook, or GPF loan management",
                "status": "missing",
                "gap": "Government provident fund entirely absent.",
                "recommendation": "Required only for government sector product.",
            },
            {
                "name": "NPS / PRAN (National Pension System)",
                "karnataka": "NPS account details for post-2004 employees; PRAN number on payslip",
                "isarva": "Not available — no PRAN field, NPS contribution calculation, or Tier-1/Tier-2 tracking",
                "status": "missing",
                "gap": "No NPS module for government-style pension.",
                "recommendation": "Add if targeting PSU/government contractors.",
            },
            {
                "name": "GIS (Group Insurance Scheme)",
                "karnataka": "GIS installment and coverage details on payslip and ESS",
                "isarva": "Not available",
                "status": "missing",
                "gap": "No group insurance scheme deduction or policy tracking.",
                "recommendation": "Government sector only.",
            },
            {
                "name": "KGID (Karnataka Govt Insurance Dept)",
                "karnataka": "KGID policy details linked to employee identity and salary deductions",
                "isarva": "Not available",
                "status": "missing",
                "gap": "State-specific insurance scheme not supported.",
                "recommendation": "Karnataka government only.",
            },
            {
                "name": "ESI (Employees State Insurance)",
                "karnataka": "Limited applicability in govt HRMS",
                "isarva": "Auto-calculated: 0.75% employee contribution when gross ≤ ₹21,000; ESI Excel export",
                "status": "full",
                "gap": "Private-sector ESI is implemented.",
                "recommendation": "Core strength for eligible establishments.",
            },
            {
                "name": "Professional Tax (PT) — Karnataka Slabs",
                "karnataka": "Karnataka govt PT deducted per state rules on govt salary (shown on payslip)",
                "isarva": "Simplified rule: ₹200/month if gross ≥ ₹25,000, else ₹0 — single flat threshold",
                "status": "partial",
                "gap": "Does NOT implement full Karnataka PT slab structure (₹0–₹200 tiered by salary bands, half-yearly filing nuances, special category exemptions).",
                "recommendation": "Implement Karnataka PT slab table (and other state PT masters) with auto-selection by employee work state.",
            },
            {
                "name": "TDS (Tax Deducted at Source / Income Tax)",
                "karnataka": "TDS deducted monthly; annual salary statement for IT return filing",
                "isarva": "TDS slab configuration UI exists in salary settings; TDS shown in analytics/comparison reports — BUT TDS is NOT auto-calculated in monthly payroll engine (recalculateEmployeePayroll only runs EPF, ESI, PT)",
                "status": "partial",
                "gap": "TDS settings exist on paper but are not wired into live monthly salary calculation. No Form 16, Form 16A, or Form 24Q generation.",
                "recommendation": "Priority fix: integrate TDS into payroll calculation using old/new tax regime, declarations (80C, HRA), and generate Form 16 annually.",
            },
            {
                "name": "Form 16 / Form 24Q (Income Tax Returns)",
                "karnataka": "Annual salary statement supports employee IT filing",
                "isarva": "Not available — no Form 16 PDF, no quarterly 24Q TDS return file",
                "status": "missing",
                "gap": "No statutory income tax return documents for employees or employer.",
                "recommendation": "Build Form 16 generator and 24Q export — high value for private sector.",
            },
            {
                "name": "HRA exemption calculation",
                "karnataka": "HRA as govt allowance component — exemption rules differ from private sector",
                "isarva": "HRA as configurable salary component; no automatic HRA tax exemption (metro/non-metro, rent declaration)",
                "status": "missing",
                "gap": "No rent declaration, metro flag, or HRA exemption in TDS computation.",
                "recommendation": "Part of TDS module — capture rent paid, city type, compute exempt HRA.",
            },
            {
                "name": "Investment declarations (80C, 80D, etc.)",
                "karnataka": "Managed through govt establishment / DDO for TDS",
                "isarva": "Not available — no employee tax declaration portal",
                "status": "missing",
                "gap": "Employees cannot submit proof of investment for TDS adjustment.",
                "recommendation": "Add tax declaration module in Attendance self-service.",
            },
            {
                "name": "Labour Welfare Fund (LWF)",
                "karnataka": "Applicable per state rules for certain establishments",
                "isarva": "Statutory component slot exists for LWF (generic); not Karnataka LWF-specific automation",
                "status": "partial",
                "gap": "Can configure manually; no Karnataka LWF rate/period automation.",
                "recommendation": "Add Karnataka LWF rules (employee + employer, annual/half-yearly).",
            },
            {
                "name": "Gratuity calculation",
                "karnataka": "Part of retirement settlement for eligible govt staff",
                "isarva": "Not auto-calculated on exit",
                "status": "missing",
                "gap": "No gratuity accrual or payment calculation (15 days salary × years of service).",
                "recommendation": "Add gratuity module for Payment of Gratuity Act compliance.",
            },
            {
                "name": "Bonus (Payment of Bonus Act)",
                "karnataka": "Festival advance / bonus as govt allowance",
                "isarva": "Can add as salary component; no statutory bonus calculation engine",
                "status": "partial",
                "gap": "No auto 8.33% bonus calculation with eligibility rules.",
                "recommendation": "Add bonus calculation module if manufacturing clients need it.",
            },
            {
                "name": "EPF ECR / compliance filing export",
                "karnataka": "N/A (uses GPF)",
                "isarva": "EPF Excel/CSV export with UAN, wages, contributions, NCP days",
                "status": "full",
                "gap": "EPF filing export is a private-sector strength.",
                "recommendation": "Keep updated with EPFO ECR format changes.",
            },
            {
                "name": "ESI contribution filing export",
                "karnataka": "N/A",
                "isarva": "ESI Excel export for monthly contribution filing",
                "status": "full",
                "gap": "ESI export available.",
                "recommendation": "None.",
            },
        ],
    },
    {
        "id": "7",
        "title": "Leave Management",
        "intro": (
            "Karnataka government leave types include Earned Leave (EL), Half Pay Leave (HPL), "
            "and other special leaves governed by service rules. ISARVA uses configurable leave "
            "types with manager/HR approval workflow."
        ),
        "features": [
            {
                "name": "Leave types",
                "karnataka": "EL, HPL, Casual Leave, Maternity, etc. per service rules; balances on ESS",
                "isarva": "Configurable leave types (LeaveTypeController); synced from Attendance to Payroll API",
                "status": "partial",
                "gap": "Flexible leave types but no govt service-rule templates (EL/HPL accrual formulas).",
                "recommendation": "Add leave accrual rule engine (monthly/yearly, carry-forward, encashment).",
            },
            {
                "name": "Leave application workflow",
                "karnataka": "Apply via ESS; approved by reporting officer / DDO",
                "isarva": "Employee applies → Manager approves → HR approves; notifications and activity log",
                "status": "full",
                "gap": "Multi-level approval workflow exists.",
                "recommendation": "None.",
            },
            {
                "name": "Leave balance visibility (employee)",
                "karnataka": "Real-time EL, HPL balances on ESS and mobile app",
                "isarva": "Leave balance fetched from Payroll API; visible in Attendance self-service",
                "status": "full",
                "gap": "Employee can view balances.",
                "recommendation": "None.",
            },
            {
                "name": "Leave encashment",
                "karnataka": "EL encashment on retirement or as per rules",
                "isarva": "Not automated — would need manual salary component adjustment",
                "status": "missing",
                "gap": "No leave encashment calculation or payout workflow.",
                "recommendation": "Add encashment module linked to leave balance and payroll.",
            },
            {
                "name": "On Other Duty (OOD) / field duty",
                "karnataka": "OOD module in Kartavya-KAAMS for field staff not marking office attendance",
                "isarva": "GPS visit check-in type exists; no formal OOD approval workflow",
                "status": "partial",
                "gap": "Field visit tracking via GPS but no OOD leave equivalent with approval.",
                "recommendation": "Add OOD/official tour request module.",
            },
            {
                "name": "Public holidays",
                "karnataka": "State gazette holidays applied centrally",
                "isarva": "Public holiday master per financial year; department-wise holiday quota config; flexible/fixed types",
                "status": "full",
                "gap": "Holiday management is comprehensive.",
                "recommendation": "None.",
            },
            {
                "name": "Compensatory off (comp-off)",
                "karnataka": "Earned against holiday/weekend work per service rules",
                "isarva": "Comp-off support in leave application flow",
                "status": "full",
                "gap": "Comp-off available.",
                "recommendation": "None.",
            },
            {
                "name": "Leave reports",
                "karnataka": "Department-wise leave monitoring dashboards",
                "isarva": "Approved/rejected leaves, LOP, monthly leave chart, daily leave PDF reports",
                "status": "full",
                "gap": "Good leave reporting for private sector.",
                "recommendation": "Add state/district hierarchy reports for large orgs.",
            },
        ],
    },
    {
        "id": "8",
        "title": "Attendance Management",
        "intro": (
            "Karnataka has deployed Kartavya-KAAMS — an AI-powered mobile attendance system with "
            "facial recognition, liveness detection, Aadhaar e-KYC, and 100–150 metre geofencing. "
            "It integrates with HRMS and SATS. ISARVA uses biometric Excel upload, TimeStation API, "
            "manual punches, and GPS geofencing."
        ),
        "features": [
            {
                "name": "Mobile attendance app",
                "karnataka": "Kartavya-KAAMS Android/iOS app; 3+ lakh registrations; state/district/taluk dashboards",
                "isarva": "GPS tracking API for mobile check-in/out; no dedicated employee attendance app in app stores",
                "status": "partial",
                "gap": "GPS API exists but no polished public mobile app like KAAMS.",
                "recommendation": "Launch branded ISARVA Attendance mobile app (Android/iOS).",
            },
            {
                "name": "Facial recognition attendance",
                "karnataka": "AI facial recognition with liveness detection to prevent photo spoofing",
                "isarva": "Not available — biometric via device Excel upload only, not face recognition",
                "status": "missing",
                "gap": "No face-based attendance capture.",
                "recommendation": "Evaluate face recognition SDK for mobile app if fraud is a concern.",
            },
            {
                "name": "Geofencing",
                "karnataka": "100–150 metre radius geofence; GIS-enabled; multiple approved locations (schools, offices, training centres)",
                "isarva": "Single office geofence (configurable lat/lng/radius ~250m); GPS visit mode for field staff",
                "status": "partial",
                "gap": "One office geofence; no multi-location geofence per employee/branch.",
                "recommendation": "Support multiple geofence zones per location/employee.",
            },
            {
                "name": "Biometric device integration",
                "karnataka": "Integrated with school/office biometric infrastructure via KAAMS",
                "isarva": "Biometric Excel import (multi-format), TimeStation API fetch, manual punch entry",
                "status": "full",
                "gap": "Strong admin-side biometric ingestion for private sector.",
                "recommendation": "Add real-time biometric device API (not just Excel).",
            },
            {
                "name": "Shift & duty roster",
                "karnataka": "School/office duty schedules",
                "isarva": "Shift master + duty roster (per employee per day); bulk week copy/clear",
                "status": "full",
                "gap": "Shift/roster management is solid.",
                "recommendation": "None.",
            },
            {
                "name": "Monthly attendance processing",
                "karnataka": "Attendance flows to HRMS for salary; SATS integration for schools",
                "isarva": "Bulk attendance generate → preview → edit → lock → Payroll API consumes locked data",
                "status": "full",
                "gap": "Lock-before-payroll is a strong control.",
                "recommendation": "Highlight as audit-friendly feature.",
            },
            {
                "name": "Late coming / early going / undertime",
                "karnataka": "Monitored via KAAMS dashboards",
                "isarva": "AttendancePolicy rules; late/early/OT/undertime calculated in AttendanceService",
                "status": "full",
                "gap": "Policy-based time deviation tracking exists.",
                "recommendation": "None.",
            },
            {
                "name": "Overtime management",
                "karnataka": "Allowance-based in govt payroll",
                "isarva": "Dedicated OT module: monthly grid, approve, lock; feeds Payroll OT/incentive module",
                "status": "full",
                "gap": "OT workflow is well implemented.",
                "recommendation": "None.",
            },
            {
                "name": "Real-time attendance dashboards (hierarchy)",
                "karnataka": "State → District → Taluk → Division monitoring; 90%+ registration tracking",
                "isarva": "Admin dashboard and GPS tracking map; no govt-style geographic hierarchy",
                "status": "partial",
                "gap": "No multi-level geographic drill-down dashboards.",
                "recommendation": "Add org-tree attendance dashboard (company → region → branch).",
            },
            {
                "name": "Employee web punch clock",
                "karnataka": "Mobile app primary; no web punch",
                "isarva": "No employee web punch; GPS mobile API or admin-processed biometric",
                "status": "partial",
                "gap": "Both rely on mobile/admin — neither has simple web clock-in.",
                "recommendation": "Optional web punch for WFH employees.",
            },
        ],
    },
    {
        "id": "9",
        "title": "Loans, Advances & Recoveries",
        "intro": (
            "Karnataka HRMS tracks GPF loans, KGID loans, festival advances, and salary recoveries "
            "visible on ESS. ISARVA handles salary advances and held salary for private sector."
        ),
        "features": [
            {
                "name": "GPF loan management",
                "karnataka": "GPF loan apply, approval, EMI deduction, balance on ESS",
                "isarva": "Not available",
                "status": "missing",
                "gap": "No provident fund loan module.",
                "recommendation": "Government sector only.",
            },
            {
                "name": "Salary advance",
                "karnataka": "Festival advance and other govt advances with recovery schedule",
                "isarva": "EmployeeAdvance module with auto-deduction at payroll finalize",
                "status": "full",
                "gap": "Salary advance with payroll recovery works.",
                "recommendation": "None.",
            },
            {
                "name": "Loan/advance visibility on ESS",
                "karnataka": "HRMS2 ESS mobile app shows loans, advances, recoveries",
                "isarva": "Employee can view payslips in Attendance; advance balance not prominently in self-service",
                "status": "partial",
                "gap": "Advance details not fully exposed in employee portal.",
                "recommendation": "Add 'My Advances' section in Attendance self-service.",
            },
            {
                "name": "Recovery schedule / EMI tracking",
                "karnataka": "Detailed recovery schedule per loan type",
                "isarva": "Advance with installment deduction; no full amortization schedule UI",
                "status": "partial",
                "gap": "Basic recovery at payroll; no loan statement or EMI calendar.",
                "recommendation": "Add advance/loan statement for employees.",
            },
        ],
    },
    {
        "id": "10",
        "title": "Retirement & Pension",
        "intro": (
            "Karnataka HRMS includes retirement management (Kellton scope) and pension slip for "
            "retired employees. ISARVA handles employee exit but not pension lifecycle."
        ),
        "features": [
            {
                "name": "Retirement management",
                "karnataka": "Retirement processing, pension commencement, service verification",
                "isarva": "Exit employee workflow with resignation date; login blocked after exit",
                "status": "partial",
                "gap": "Exit exists; no retirement age trigger, superannuation workflow, or pension initiation.",
                "recommendation": "Add retirement due report and superannuation checklist.",
            },
            {
                "name": "Pension slip",
                "karnataka": "Pension statement for retired employees on portal",
                "isarva": "Not available",
                "status": "missing",
                "gap": "No post-retirement pension payment module.",
                "recommendation": "Government sector only.",
            },
            {
                "name": "Gratuity / retiral benefits",
                "karnataka": "Calculated as part of govt retirement settlement",
                "isarva": "Not auto-calculated",
                "status": "missing",
                "gap": "See Tax module — gratuity not implemented.",
                "recommendation": "Implement for private sector Payment of Gratuity Act.",
            },
            {
                "name": "Leave encashment on retirement",
                "karnataka": "EL encashment processed in final settlement",
                "isarva": "Not automated",
                "status": "missing",
                "gap": "No retirement leave encashment workflow.",
                "recommendation": "Link leave balance to FFS on exit.",
            },
        ],
    },
    {
        "id": "11",
        "title": "Employee Self-Service (ESS)",
        "intro": (
            "Karnataka HRMS2 ESS (mobile app + hrmsess.karnataka.gov.in) gives employees direct "
            "access to payslips, leave, service book, loans, and transfer documents without "
            "going through DDO. ISARVA routes employees to Attendance portal for self-service."
        ),
        "features": [
            {
                "name": "ESS portal / mobile app",
                "karnataka": "HRMS2 ESS app (Google Play); payslips, leave, loans, service book, transfer docs",
                "isarva": "Attendance module as self-service portal: leave, payslips, profile, data change request",
                "status": "partial",
                "gap": "Self-service exists but split across Hub + Attendance; no unified ESS mobile app.",
                "recommendation": "Package Attendance as 'ISARVA ESS' with unified branding.",
            },
            {
                "name": "Payslip download",
                "karnataka": "Current and past months via ESS",
                "isarva": "View and download PDF via Attendance /payroll API",
                "status": "full",
                "gap": "Implemented.",
                "recommendation": "None.",
            },
            {
                "name": "Service book view",
                "karnataka": "Full e-SR on ESS",
                "isarva": "Not available",
                "status": "missing",
                "gap": "No career history view for employees.",
                "recommendation": "Add employee timeline (joining, promotions, transfers, trainings).",
            },
            {
                "name": "Transfer document access",
                "karnataka": "Transfer orders on ESS",
                "isarva": "Not available",
                "status": "missing",
                "gap": "No transfer order repository for employees.",
                "recommendation": "Depends on transfer module.",
            },
            {
                "name": "Tax declaration submission",
                "karnataka": "Via establishment / integrated in salary processing",
                "isarva": "Not available",
                "status": "missing",
                "gap": "See TDS gaps.",
                "recommendation": "High priority for private sector ESS.",
            },
        ],
    },
    {
        "id": "12",
        "title": "Reports, Analytics & Integrations",
        "intro": (
            "Karnataka HRMS 2.0 targets 200+ API integrations with 45+ departments and data "
            "analytics dashboards. ISARVA integrates Payroll ↔ Attendance via REST APIs and webhooks."
        ),
        "features": [
            {
                "name": "Operational reports",
                "karnataka": "Salary bills, establishment reports, leave statements, attendance summaries",
                "isarva": "Payroll reports, OT/incentive reports, leave reports, analytical payroll comparison",
                "status": "partial",
                "gap": "Good private-sector reports; no govt establishment/registrar reports.",
                "recommendation": "Add custom report builder for enterprise clients.",
            },
            {
                "name": "Data analytics / BI dashboards",
                "karnataka": "Analytical reports with data analytics (Kellton HRMS 2.0 scope)",
                "isarva": "Dashboard widgets, payroll comparison charts, leave analytics",
                "status": "partial",
                "gap": "Basic analytics; no advanced BI or predictive workforce analytics.",
                "recommendation": "Integrate Metabase/Power BI or build workforce cost dashboard.",
            },
            {
                "name": "Cross-department API integrations",
                "karnataka": "200+ APIs with 45+ Karnataka departments + central govt",
                "isarva": "Payroll ↔ Attendance APIs; Hub SSO; CRM/POSH user sync; webhook-based employee sync",
                "status": "partial",
                "gap": "Strong internal integration; no external govt system APIs (treasury, IFMS, etc.).",
                "recommendation": "Document public API for third-party integrations.",
            },
            {
                "name": "Audit trail / activity log",
                "karnataka": "Government audit requirements",
                "isarva": "Spatie Activity Log in Attendance; ActivityLogService in Payroll for payroll actions",
                "status": "full",
                "gap": "Audit logging present.",
                "recommendation": "None.",
            },
            {
                "name": "Compliance (DPDP / data privacy)",
                "karnataka": "Government data governance",
                "isarva": "DPDP consent flow, data change request module, POSH compliance (separate app)",
                "status": "full",
                "gap": "Modern privacy compliance for private sector.",
                "recommendation": "Strength for corporate clients.",
            },
        ],
    },
]

PRIORITY_RECOMMENDATIONS = [
    ("P1 — Critical for private sector", [
        "Wire TDS into monthly payroll calculation (currently settings-only)",
        "Implement Karnataka (and multi-state) Professional Tax slab engine",
        "Generate Form 16 and annual salary statement for employees",
        "Add employee tax declaration module (80C, HRA, regime selection)",
    ]),
    ("P2 — High value enhancements", [
        "Launch ISARVA Attendance mobile app (GPS + optional face recognition)",
        "Multi-location geofence support for branch-based attendance",
        "Gratuity calculation on employee exit",
        "Leave encashment automation",
        "Employee 'My Advances & Loans' in self-service portal",
    ]),
    ("P3 — Government sector only (if entering that market)", [
        "GPF / NPS / GIS / KGID modules",
        "Electronic Service Book (e-SR)",
        "Cadre management, transfer orders, DDO workflow",
        "Recruitment and APR modules",
        "Pension management post-retirement",
    ]),
]

SUMMARY_COUNTS = {"full": 0, "partial": 0, "missing": 0, "na": 0}
for mod in MODULES:
    for f in mod["features"]:
        SUMMARY_COUNTS[f["status"]] = SUMMARY_COUNTS.get(f["status"], 0) + 1


# ---------------------------------------------------------------------------
# Markdown generator
# ---------------------------------------------------------------------------

def build_markdown() -> str:
    lines = [
        f"# {META['title']}",
        "",
        f"**Prepared for:** {META['prepared_for']}  ",
        f"**Benchmark:** {META['benchmark']}  ",
        f"**Version:** {META['version']}  ",
        f"**Date:** {META['date']}",
        "",
        "---",
        "",
        "## Document Purpose",
        "",
        "This document compares ISARVA HRMS (Payroll and Attendance modules) against "
        "Karnataka State Government HRMS 2.0. It is written in plain language so that HR managers, "
        "payroll administrators, and business stakeholders can understand what ISARVA already "
        "does well, where gaps exist, and what would need to be built.",
        "",
        "## Status Legend",
        "",
        "| Status | Meaning |",
        "|--------|---------|",
        "| **FULL** | Feature is available and functional in ISARVA |",
        "| **PARTIAL** | Feature exists but is incomplete or simplified |",
        "| **GAP** | Feature is not available in ISARVA |",
        "| **N/A** | Not applicable — different target market |",
        "",
        "## Summary Scorecard",
        "",
        f"| Status | Count |",
        f"|--------|-------|",
        f"| Full / Available | {SUMMARY_COUNTS['full']} |",
        f"| Partial | {SUMMARY_COUNTS['partial']} |",
        f"| Gap (Not Available) | {SUMMARY_COUNTS['missing']} |",
        f"| Not Applicable | {SUMMARY_COUNTS['na']} |",
        "",
        "---",
        "",
    ]

    for mod in MODULES:
        lines += [f"## Module {mod['id']}: {mod['title']}", ""]
        if mod.get("intro"):
            for para in mod["intro"].strip().split("\n\n"):
                lines += [para, ""]
        lines += [
            "| Feature | Karnataka HRMS 2.0 | ISARVA HRMS | Status | Gap / Notes |",
            "|---------|-------------------|-------------|--------|-------------|",
        ]
        for f in mod["features"]:
            label, code, _ = STATUS[f["status"]]
            lines.append(
                f"| {f['name']} | {f['karnataka']} | {f['isarva']} | **{code}** | {f['gap']} |"
            )
        lines += ["", "### Recommendations", ""]
        recs = {f["name"]: f["recommendation"] for f in mod["features"] if f["recommendation"] != "None." and f["recommendation"] != "None"}
        if recs:
            for name, rec in recs.items():
                lines.append(f"- **{name}:** {rec}")
        else:
            lines.append("- No critical actions for this module.")
        lines += ["", "---", ""]

    lines += ["## Priority Roadmap", ""]
    for title, items in PRIORITY_RECOMMENDATIONS:
        lines += [f"### {title}", ""]
        for item in items:
            lines.append(f"1. {item}")
        lines.append("")

    lines += [
        "## Important Disclaimer",
        "",
        "Karnataka HRMS 2.0 is a government system designed for 6+ lakh state employees under "
        "government pay rules. ISARVA HRMS is designed for private companies under EPF/ESI/CTC "
        "payroll rules. A 'GAP' does not always mean ISARVA is deficient — it often means the "
        "feature is government-specific and not needed for corporate clients.",
        "",
        f"*Generated automatically on {META['date']} by generate_gap_analysis.py*",
    ]
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Shared table helpers
# ---------------------------------------------------------------------------

def _pdf_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _pdf_para(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(_pdf_escape(text), style)


def _feature_table_rows(module: dict) -> list[list]:
    rows = [["Feature", "Karnataka HRMS 2.0", "ISARVA HRMS", "Status", "Gap / Notes"]]
    for f in module["features"]:
        _, code, _ = STATUS[f["status"]]
        rows.append([f["name"], f["karnataka"], f["isarva"], code, f["gap"]])
    return rows


DOC_PURPOSE = (
    "This document compares ISARVA HRMS (Payroll and Attendance modules) against "
    "Karnataka State Government HRMS 2.0. It is written in plain language so that HR managers, "
    "payroll administrators, and business stakeholders can understand what ISARVA already "
    "does well, where gaps exist, and what would need to be built."
)

DISCLAIMER = (
    "Karnataka HRMS 2.0 is a government system designed for 6+ lakh state employees under "
    "government pay rules. ISARVA HRMS is designed for private companies under EPF/ESI/CTC "
    "payroll rules. A 'GAP' does not always mean ISARVA is deficient — it often means the "
    "feature is government-specific and not needed for corporate clients."
)


# ---------------------------------------------------------------------------
# DOCX generator
# ---------------------------------------------------------------------------

def _set_doc_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _docx_fill_table_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def build_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    _set_doc_styles(doc)

    title = doc.add_heading(META["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        f"Prepared for: {META['prepared_for']}",
        f"Benchmark: {META['benchmark']}",
        f"Version: {META['version']} | Date: {META['date']}",
    ]:
        run = meta.add_run(line + "\n")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_heading("Document Purpose", level=1)
    doc.add_paragraph(DOC_PURPOSE)

    doc.add_heading("Status Legend", level=1)
    legend_table = doc.add_table(rows=5, cols=2)
    legend_table.style = "Table Grid"
    legend_rows = [
        ("Status", "Meaning"),
        ("FULL", "Feature is available and functional in ISARVA"),
        ("PARTIAL", "Feature exists but is incomplete or simplified"),
        ("GAP", "Feature is not available in ISARVA"),
        ("N/A", "Not applicable — different target market"),
    ]
    for i, (a, b) in enumerate(legend_rows):
        _docx_fill_table_cell(legend_table.rows[i].cells[0], a, bold=(i == 0))
        _docx_fill_table_cell(legend_table.rows[i].cells[1], b, bold=(i == 0))

    doc.add_heading("Summary Scorecard", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Status", "Count"),
        ("Full / Available", str(SUMMARY_COUNTS["full"])),
        ("Partial", str(SUMMARY_COUNTS["partial"])),
        ("Gap (Not Available)", str(SUMMARY_COUNTS["missing"])),
        ("Not Applicable", str(SUMMARY_COUNTS["na"])),
    ]
    for i, (a, b) in enumerate(rows_data):
        _docx_fill_table_cell(table.rows[i].cells[0], a, bold=(i == 0))
        _docx_fill_table_cell(table.rows[i].cells[1], b, bold=(i == 0))

    for mod in MODULES:
        doc.add_page_break()
        doc.add_heading(f"Module {mod['id']}: {mod['title']}", level=1)
        if mod.get("intro"):
            doc.add_paragraph(mod["intro"])

        feature_rows = _feature_table_rows(mod)
        feat_table = doc.add_table(rows=len(feature_rows), cols=5)
        feat_table.style = "Table Grid"
        col_widths = [Inches(1.2), Inches(1.8), Inches(1.8), Inches(0.7), Inches(1.8)]
        for row_idx, row_data in enumerate(feature_rows):
            for col_idx, cell_text in enumerate(row_data):
                _docx_fill_table_cell(
                    feat_table.rows[row_idx].cells[col_idx],
                    cell_text,
                    bold=(row_idx == 0),
                )
                feat_table.rows[row_idx].cells[col_idx].width = col_widths[col_idx]

        doc.add_heading("Recommendations", level=2)
        recs = {
            f["name"]: f["recommendation"]
            for f in mod["features"]
            if f["recommendation"] not in ("None.", "None")
        }
        if recs:
            rec_table = doc.add_table(rows=len(recs) + 1, cols=2)
            rec_table.style = "Table Grid"
            _docx_fill_table_cell(rec_table.rows[0].cells[0], "Feature", bold=True)
            _docx_fill_table_cell(rec_table.rows[0].cells[1], "Recommendation", bold=True)
            for i, (name, rec) in enumerate(recs.items(), start=1):
                _docx_fill_table_cell(rec_table.rows[i].cells[0], name)
                _docx_fill_table_cell(rec_table.rows[i].cells[1], rec)
        else:
            doc.add_paragraph("No critical actions for this module.")

    doc.add_page_break()
    doc.add_heading("Priority Roadmap", level=1)
    for title, items in PRIORITY_RECOMMENDATIONS:
        doc.add_heading(title, level=2)
        roadmap_table = doc.add_table(rows=len(items) + 1, cols=2)
        roadmap_table.style = "Table Grid"
        _docx_fill_table_cell(roadmap_table.rows[0].cells[0], "#", bold=True)
        _docx_fill_table_cell(roadmap_table.rows[0].cells[1], "Action Item", bold=True)
        for i, item in enumerate(items, start=1):
            _docx_fill_table_cell(roadmap_table.rows[i].cells[0], str(i))
            _docx_fill_table_cell(roadmap_table.rows[i].cells[1], item)

    doc.add_heading("Important Disclaimer", level=1)
    doc.add_paragraph(DISCLAIMER)
    doc.add_paragraph(f"Generated automatically on {META['date']} by generate_gap_analysis.py")

    doc.save(str(docx_path))
    print(f"DOCX written: {docx_path}")


# ---------------------------------------------------------------------------
# PDF generator (ReportLab)
# ---------------------------------------------------------------------------

def _pdf_table_style(header_rows: int = 1) -> TableStyle:
    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, header_rows - 1), colors.HexColor("#1A5276")),
        ("TEXTCOLOR", (0, 0), (-1, header_rows - 1), colors.white),
        ("FONTNAME", (0, 0), (-1, header_rows - 1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, header_rows), (-1, -1), [colors.white, colors.HexColor("#F8F9F9")]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D5D8DC")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
    ])


def build_pdf(pdf_path: Path) -> None:
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title=META["title"],
        author="ISARVA HRMS",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CoverTitle", fontSize=22, leading=28, alignment=TA_CENTER,
        textColor=colors.HexColor("#1A5276"), spaceAfter=12, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="CoverSub", fontSize=12, leading=16, alignment=TA_CENTER,
        textColor=colors.HexColor("#566573"), spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="ModuleTitle", fontSize=14, leading=18, spaceBefore=14, spaceAfter=8,
        textColor=colors.HexColor("#1A5276"), fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="SectionTitle", fontSize=11, leading=14, spaceBefore=8, spaceAfter=6,
        textColor=colors.HexColor("#2C3E50"), fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="BodyJustify", fontSize=9, leading=13, alignment=TA_JUSTIFY, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="TableCell", fontSize=7, leading=9, alignment=TA_LEFT, spaceAfter=0,
    ))
    styles.add(ParagraphStyle(
        name="TableHeader", fontSize=7, leading=9, alignment=TA_LEFT, spaceAfter=0,
        textColor=colors.white, fontName="Helvetica-Bold",
    ))

    col_widths = [2.2 * cm, 4.3 * cm, 4.3 * cm, 1.4 * cm, 4.3 * cm]
    cell_style = styles["TableCell"]
    header_style = styles["TableHeader"]

    story = []

    # Cover
    story.append(Spacer(1, 2 * cm))
    story.append(Paragraph(META["title"], styles["CoverTitle"]))
    story.append(Paragraph(META["subtitle"], styles["CoverSub"]))
    story.append(Spacer(1, 0.5 * cm))
    for line in [
        f"<b>Prepared for:</b> {META['prepared_for']}",
        f"<b>Benchmark:</b> {META['benchmark']}",
        f"<b>Version:</b> {META['version']} &nbsp;|&nbsp; <b>Date:</b> {META['date']}",
    ]:
        story.append(Paragraph(line, styles["CoverSub"]))
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(_pdf_escape(DOC_PURPOSE), styles["BodyJustify"]))
    story.append(PageBreak())

    # Status Legend
    story.append(Paragraph("Status Legend", styles["ModuleTitle"]))
    legend_data = [
        [_pdf_para("Status", header_style), _pdf_para("Meaning", header_style)],
        [_pdf_para("FULL", cell_style), _pdf_para("Feature is available and functional in ISARVA", cell_style)],
        [_pdf_para("PARTIAL", cell_style), _pdf_para("Feature exists but is incomplete or simplified", cell_style)],
        [_pdf_para("GAP", cell_style), _pdf_para("Feature is not available in ISARVA", cell_style)],
        [_pdf_para("N/A", cell_style), _pdf_para("Not applicable — different target market", cell_style)],
    ]
    lt = Table(legend_data, colWidths=[3 * cm, 14 * cm])
    lt.setStyle(_pdf_table_style())
    story.append(lt)
    story.append(Spacer(1, 0.4 * cm))

    # Scorecard
    story.append(Paragraph("Summary Scorecard", styles["ModuleTitle"]))
    score_data = [
        [_pdf_para("Status", header_style), _pdf_para("Count", header_style)],
        [_pdf_para("Full / Available", cell_style), _pdf_para(str(SUMMARY_COUNTS["full"]), cell_style)],
        [_pdf_para("Partial", cell_style), _pdf_para(str(SUMMARY_COUNTS["partial"]), cell_style)],
        [_pdf_para("Gap (Not Available)", cell_style), _pdf_para(str(SUMMARY_COUNTS["missing"]), cell_style)],
        [_pdf_para("Not Applicable", cell_style), _pdf_para(str(SUMMARY_COUNTS["na"]), cell_style)],
    ]
    t = Table(score_data, colWidths=[8 * cm, 4 * cm])
    t.setStyle(_pdf_table_style())
    story.append(t)
    story.append(PageBreak())

    for mod in MODULES:
        story.append(Paragraph(f"Module {mod['id']}: {mod['title']}", styles["ModuleTitle"]))
        if mod.get("intro"):
            for para in mod["intro"].strip().split("\n\n"):
                story.append(Paragraph(_pdf_escape(para.replace("\n", " ")), styles["BodyJustify"]))
        story.append(Spacer(1, 0.2 * cm))

        rows = _feature_table_rows(mod)
        table_data = []
        for r_idx, row in enumerate(rows):
            style = header_style if r_idx == 0 else cell_style
            table_data.append([_pdf_para(cell, style) for cell in row])

        feat_table = Table(table_data, colWidths=col_widths, repeatRows=1)
        feat_table.setStyle(_pdf_table_style())
        story.append(feat_table)
        story.append(Spacer(1, 0.3 * cm))

        story.append(Paragraph("Recommendations", styles["SectionTitle"]))
        recs = {
            f["name"]: f["recommendation"]
            for f in mod["features"]
            if f["recommendation"] not in ("None.", "None")
        }
        if recs:
            rec_data = [
                [_pdf_para("Feature", header_style), _pdf_para("Recommendation", header_style)],
            ]
            for name, rec in recs.items():
                rec_data.append([_pdf_para(name, cell_style), _pdf_para(rec, cell_style)])
            rec_table = Table(rec_data, colWidths=[4.5 * cm, 12 * cm], repeatRows=1)
            rec_table.setStyle(_pdf_table_style())
            story.append(rec_table)
        else:
            story.append(Paragraph("No critical actions for this module.", styles["BodyJustify"]))

        story.append(PageBreak())

    # Priority roadmap
    story.append(Paragraph("Priority Roadmap", styles["ModuleTitle"]))
    for title, items in PRIORITY_RECOMMENDATIONS:
        story.append(Paragraph(title, styles["SectionTitle"]))
        roadmap_data = [
            [_pdf_para("#", header_style), _pdf_para("Action Item", header_style)],
        ]
        for i, item in enumerate(items, start=1):
            roadmap_data.append([_pdf_para(str(i), cell_style), _pdf_para(item, cell_style)])
        rt = Table(roadmap_data, colWidths=[1 * cm, 15.5 * cm], repeatRows=1)
        rt.setStyle(_pdf_table_style())
        story.append(rt)
        story.append(Spacer(1, 0.2 * cm))

    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("Important Disclaimer", styles["SectionTitle"]))
    story.append(Paragraph(_pdf_escape(DISCLAIMER), styles["BodyJustify"]))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(
        f"<i>Generated automatically on {META['date']} by generate_gap_analysis.py</i>",
        styles["BodyJustify"],
    ))

    doc.build(story)
    print(f"PDF written: {pdf_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    md_path = OUTPUT_DIR / "HRMS_GAP_ANALYSIS_KARNATAKA_2.0.md"
    pdf_path = OUTPUT_DIR / "HRMS_GAP_ANALYSIS_KARNATAKA_2.0.pdf"
    docx_path = OUTPUT_DIR / "HRMS_GAP_ANALYSIS_KARNATAKA_2.0.docx"

    md_content = build_markdown()
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Markdown written: {md_path}")

    build_docx(md_path, docx_path)
    build_pdf(pdf_path)

    print("\nDone. Generated files:")
    for p in (md_path, pdf_path, docx_path):
        size_kb = p.stat().st_size / 1024
        print(f"  - {p} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
